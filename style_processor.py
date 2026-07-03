"""
style_processor.py
Heurísticas validadas con 25 pares de documentos reales (2,760 cambios).
"""

import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph as DocxParagraph

# ── Constantes ────────────────────────────────────────────────────────────────
HEADING_MAX_WORDS       = 7    # Para Heading 1
HEADING_MAX_WORDS_H2    = 20   # Para Heading 2 (SSCP tienen sub-secciones largas)
HEADING_MAX_CHARS       = 120
HEADING_CONFIDENCE_HIGH = 0.90
HEADING_CONFIDENCE_MED  = 0.65
HEADING_CONFIDENCE_LOW  = 0.40

# Estilos Heading que pueden estar mal aplicados (sin bold explícito = posible Normal)
DEMOTABLE_HEADINGS = {
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9", "Title"
}
HEADING_TARGETS = {"Heading 1", "Heading 2", "Heading 3"}

# Estilos de tabla — destinos válidos para contenido dentro de tablas
TABLE_TARGET_STYLES = {"Table Paragraph", "Table Header", "Table Header 2"}

# Estilos que deben sugerirse como Table Paragraph cuando están en tabla
NON_TABLE_BODY_STYLES = {
    "Normal", "Default", "Body Text", "No Spacing",
    "Numbered List 1", "List Bullet 2", "Style1",
}

# Estilos no en el XML de Oxygen pero válidos como destino — se dejan sin cambio
TOLERATED_STYLES = {
    "Normal",           # Estilo base de Word
    "Default",          # Variante de Normal (cuando no se convierte a heading)
    "Table Paragraph",  # Contenido de celdas — reconocido por Oxygen como <p>/<entry>
    "Table Header",     # Cabecera de tabla
    "Table Header 2",   # Variante de cabecera
    "Quote",            # Citas textuales
    "No Spacing",       # Estilo de Word sin espaciado
    "footnote text",    # Texto de nota al pie
    "toc 1", "toc 2", "toc 3",  # Tabla de contenidos
    "annotation text",  # Anotaciones
    "Normal (Web)",     # Variante web de Normal
    "Caption",          # Pies de figura/tabla
    "List Bullet 2",    # Variante de lista
    "paragraph",        # Estilo personalizado
    "pf0", "pf1",       # FrameMaker — se dejan si no se detectan como heading
}

STATUS_LABELS = {
    "no_change": "✅ Sin cambio",
    "auto":      "🔵 Automático",
    "review":    "🟡 Revisar",
    "unknown":   "⚠️ Sin mapeo",
}


# ── Detección de negrita ──────────────────────────────────────────────────────

def _is_bold(para) -> bool:
    """
    True si algún run con texto tiene bold EXPLÍCITO (run.bold is True).
    Usamos bold explícito para is_misapplied_heading porque los headings
    mal aplicados (body text con estilo Heading) nunca tienen bold explícito.
    """
    return any(run.bold is True for run in para.runs if run.text.strip())


def _word_count(para) -> int:
    return len(para.text.split())


# ── Contexto de tabla ─────────────────────────────────────────────────────────

def is_in_table(para) -> bool:
    """True si el párrafo está dentro de una celda de tabla."""
    parent = para._element.getparent()
    while parent is not None:
        tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
        if tag == 'tbl':  return True
        if tag == 'body': return False
        parent = parent.getparent()
    return False


# ── Detección de headings ─────────────────────────────────────────────────────

# Documentos tipo FAQ usan preguntas completas en negrita como Heading 1.
# Estas suelen ser más largas que los títulos cortos tipo "WARNINGS".
HEADING_MAX_WORDS_QUESTION = 15

def is_likely_heading(para) -> tuple:
    """
    Evalúa si un párrafo parece un Heading 1 (título de sección / topic boundary).
    Criterio base: ≤ 7 palabras, bold y/o mayúsculas.
    Criterio adicional para FAQ: preguntas en negrita (termina en "?"),
    con límite más amplio de hasta 15 palabras — validado con documentos
    IFP tipo Zimmer donde los headings son preguntas completas.
    """
    text = para.text.strip()
    if not text:
        return False, 0.0, "párrafo vacío"

    wc   = _word_count(para)
    bold = _is_bold(para)

    # Caso especial: pregunta completa en negrita (formato FAQ)
    if bold and text.endswith("?") and wc <= HEADING_MAX_WORDS_QUESTION:
        return True, HEADING_CONFIDENCE_HIGH, "negrita + pregunta (formato FAQ)"

    if wc > HEADING_MAX_WORDS or len(text) > HEADING_MAX_CHARS:
        return False, 0.0, f"demasiado largo ({wc} palabras)"

    upper = text.isupper()
    title = text.istitle()

    if bold and (upper or wc <= 3):
        return True, HEADING_CONFIDENCE_HIGH, "negrita + mayúsculas/corto"
    if bold and title:
        return True, HEADING_CONFIDENCE_MED, "negrita + capitalizado"
    if bold:
        return True, HEADING_CONFIDENCE_MED, "negrita"
    if upper and wc <= 5:
        return True, HEADING_CONFIDENCE_LOW, "mayúsculas (sin negrita)"

    return False, 0.0, "no cumple criterios de heading"


def is_likely_subheading(para) -> tuple:
    """
    Evalúa si un párrafo parece un Heading 2 (sub-sección dentro de un topic).
    Criterio más amplio que H1: hasta 20 palabras, bold explícito requerido.
    Validado con 109 casos List Paragraph → Heading 2 en documentos SSCP.
    """
    text = para.text.strip()
    if not text:
        return False, 0.0, "párrafo vacío"

    bold = _is_bold(para)
    if not bold:
        return False, 0.0, "sin negrita explícita"

    wc = _word_count(para)
    if wc > HEADING_MAX_WORDS_H2:
        return False, 0.0, f"demasiado largo para sub-heading ({wc} palabras)"

    if wc <= 7:
        return True, HEADING_CONFIDENCE_HIGH, f"negrita ({wc} palabras)"
    return True, HEADING_CONFIDENCE_MED, f"negrita (largo: {wc} palabras)"


def is_misapplied_heading(para) -> tuple:
    """
    Detecta headings mal aplicados: tienen estilo Heading pero sin bold explícito.
    Validado: 0% de los 842 casos Heading→Normal tenían bold explícito.
    """
    text  = para.text.strip()
    bold  = _is_bold(para)
    empty = not text

    if bold:
        return False, 0.0, "tiene negrita — heading correcto"
    if empty:
        return True, 0.85, "heading vacío sin negrita"
    wc = _word_count(para)
    if wc > 10:
        return True, 0.90, f"heading sin negrita y largo ({wc} palabras)"
    return True, 0.65, f"heading sin negrita ({wc} palabras)"


# ── Traversal ─────────────────────────────────────────────────────────────────

def count_table_body_paragraphs(file_path) -> dict:
    """
    Cuenta párrafos con estilos de cuerpo dentro de tablas.
    Usado por la UI para sugerir si activar la conversión a Table Paragraph.
    Devuelve {"total_in_table": N, "normal_in_table": M, "suggest_conversion": bool}
    """
    try:
        doc = Document(file_path)
    except Exception:
        return {"total_in_table": 0, "normal_in_table": 0, "suggest_conversion": False}

    total = normal = 0
    for para in _get_all_paras(doc):
        if is_in_table(para):
            total += 1
            if (para.style and para.style.name or "Normal") in NON_TABLE_BODY_STYLES:
                normal += 1

    suggest = False  # Siempre OFF — el usuario decide según el tipo de documento
    return {"total_in_table": total, "normal_in_table": normal, "suggest_conversion": suggest}


def _get_all_paras(doc):
    """Párrafos en orden de documento, incluyendo dentro de tablas."""
    paras      = []
    para_index = {p._element: p for p in doc.paragraphs}

    def traverse(element):
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                paras.append(para_index.get(child) or DocxParagraph(child, doc))
            elif tag in ('tbl', 'tr', 'tc', 'sdt', 'sdtContent', 'body'):
                traverse(child)

    traverse(doc.element.body)
    return paras


# ── Análisis principal ────────────────────────────────────────────────────────

# ── Coincidencia por similitud de nombre (closest match) ──────────────────────
# Inspirado en el algoritmo equivalente de DocStyler (Java): cuando un estilo
# no tiene regla guardada ni heurística aplicable, se busca el estilo aprobado
# más parecido por nombre, combinando contención de subcadena y solapamiento
# de palabras. Es un apoyo adicional — nunca sustituye a una regla confirmada
# con datos reales, y siempre queda en estado "revisar", nunca automático.

def _normalize_style_name(name: str) -> str:
    return re.sub(r"[\s_-]+", "", name.lower())


def _split_style_words(name: str) -> set:
    parts = re.split(r"[\s_-]+", name)
    return {p.lower() for p in parts if p}


def find_closest_style(style_name: str, known_style_names, min_score: float = 0.3):
    """
    Busca el estilo aprobado (de Oxygen) más parecido a un nombre desconocido.

    Devuelve (estilo_sugerido_o_None, score_0_a_1, motivo_string).
    """
    normalized    = _normalize_style_name(style_name)
    source_words  = _split_style_words(style_name)

    best_match = None
    best_score = 0.0

    for approved in known_style_names:
        approved_norm = _normalize_style_name(approved)

        # Coincidencia exacta salvo mayúsculas/espacios — muy alta confianza
        if normalized == approved_norm:
            return approved, 0.95, "coincide salvo mayúsculas/espacios"

        score = 0.0
        if normalized in approved_norm or approved_norm in normalized:
            score = min(len(normalized), len(approved_norm)) / max(len(normalized), len(approved_norm))

        approved_words = _split_style_words(approved)
        overlap = source_words & approved_words
        if overlap:
            word_score = len(overlap) / max(len(source_words), len(approved_words))
            score = max(score, word_score)

        if score > best_score:
            best_score = score
            best_match = approved

    if best_match and best_score >= min_score:
        return best_match, best_score, "similitud de nombre"
    return None, 0.0, "sin coincidencia razonable"


# ── Validación de jerarquía de headings ───────────────────────────────────────
# Detecta saltos de nivel (p.ej. Heading 3 justo después de Heading 1, sin
# pasar por Heading 2), que generan una estructura de topics/subtopics
# incorrecta al convertir a DITA. Es informativo — no modifica nada.

HEADING_LEVEL_MAP = {
    "title": 1, "heading 1": 1,
    "subtitle": 2, "heading 2": 2,
    "heading 3": 3, "heading 4": 4, "heading 5": 5,
    "heading 6": 6, "heading 7": 7, "heading 8": 8, "heading 9": 9,
}


def check_heading_hierarchy(paragraphs) -> list:
    """
    Revisa la secuencia de estilos sugeridos (analyze_document) y detecta
    saltos de nivel en la jerarquía de headings.

    paragraphs: lista de dicts con al menos 'idx', 'text', 'suggested_style'

    Devuelve una lista de dicts: {idx, text, from_level, to_level, message}
    """
    issues = []
    prev_level = 0

    for p in paragraphs:
        level = HEADING_LEVEL_MAP.get(p["suggested_style"].lower())
        if level is None:
            continue

        if prev_level > 0 and level > prev_level + 1:
            issues.append({
                "idx": p["idx"],
                "text": p["text"],
                "from_level": prev_level,
                "to_level": level,
                "message": (
                    f"'{p['suggested_style']}' aparece justo después de un "
                    f"heading de nivel {prev_level} — se esperaba nivel "
                    f"{prev_level + 1} antes de nivel {level}"
                ),
            })
        prev_level = level

    return issues


# ── Normalización de nombres de estilo por idioma ─────────────────────────────
# Word guarda dos identificadores distintos por estilo:
#   · style_id  → invariante, siempre en inglés (ej. "Heading1")
#   · name      → nombre visible, SÍ se traduce según el idioma de la
#                 interfaz de Word (ej. "Título 1" en español)
#
# Nuestras reglas, heurísticas y el XML de OxygenAuthor comparan por NOMBRE
# VISIBLE. Un documento creado o editado en Word en español puede tener sus
# estilos estándar con nombre en español, lo que los hace irreconocibles
# aunque conceptualmente sean los mismos estilos. Esta función los detecta
# por su style_id (invariante) y renombra la definición al nombre canónico
# en inglés, para que tanto Style Mapper como OxygenAuthor los reconozcan.
STANDARD_STYLE_ID_TO_ENGLISH_NAME = {
    "Normal": "Normal",
    "Heading1": "Heading 1", "Heading2": "Heading 2", "Heading3": "Heading 3",
    "Heading4": "Heading 4", "Heading5": "Heading 5", "Heading6": "Heading 6",
    "Heading7": "Heading 7", "Heading8": "Heading 8", "Heading9": "Heading 9",
    "Title": "Title", "Subtitle": "Subtitle",
    "BodyText": "Body Text", "BodyText2": "Body Text 2", "BodyText3": "Body Text 3",
    "BodyTextIndent": "Body Text Indent",
    "Quote": "Quote", "IntenseQuote": "Intense Quote",
    "Caption": "Caption",
    "ListParagraph": "List Paragraph",
    "NoSpacing": "No Spacing",
}


def normalize_localized_style_names(doc) -> list:
    """
    Renombra las definiciones de estilo cuyo style_id coincide con un estilo
    estándar de Word pero cuyo nombre visible está en otro idioma.

    Debe ejecutarse UNA VEZ al abrir el documento, antes de analizar o
    aplicar cualquier cambio — tanto analyze_document como
    apply_paragraph_decisions dependen de que ya se haya normalizado.

    Devuelve una lista de tuplas (nombre_original, nombre_nuevo, style_id)
    con los renombrados realizados, para poder informar al usuario.
    """
    renamed = []
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        canonical = STANDARD_STYLE_ID_TO_ENGLISH_NAME.get(style.style_id)
        if canonical and style.name != canonical:
            old_name = style.name
            try:
                style.name = canonical
                renamed.append((old_name, canonical, style.style_id))
            except Exception:
                pass  # Si no se puede renombrar, se deja como está
    return renamed


def normalize_document_file(input_path, output_path) -> list:
    """
    Abre un .docx, normaliza los nombres de estilo localizados y lo guarda
    en output_path. Devuelve la lista de renombrados realizados.
    Si no hay nada que renombrar, el archivo de salida es una copia idéntica.
    """
    doc = Document(input_path)
    renamed = normalize_localized_style_names(doc)
    doc.save(output_path)
    return renamed


def analyze_document(file_path, rules_dict, known_style_names, apply_table_context=False):
    """
    Analiza todos los párrafos y devuelve lista de dicts para la revisión.

    Orden de prioridad en la clasificación:

    0. ¿Está en tabla y tiene un estilo de body text inapropiado?
       → Sugerir Table Paragraph (auto si Normal, review si otros)

    1. ¿Es un heading conocido SIN bold explícito?
       → Sugerir Normal para revisión

    2. ¿Es List Paragraph con bold explícito?
       → Sugerir Heading 2 para revisión

    3. ¿Está en known_style_names?
       → Sin cambio

    4. ¿Tiene regla guardada?
       → Aplicar regla (con heurística si target es Heading)

    5. ¿Sin regla pero parece heading?
       → Sugerir Heading 1 para revisión

    6. ¿Estilo tolerable?
       → Sin cambio

    7. Sin mapeo
    """
    try:
        doc = Document(file_path)
    except PackageNotFoundError:
        raise ValueError("El archivo no es un .docx válido.")

    all_paras = _get_all_paras(doc)
    results   = []

    for i, para in enumerate(all_paras):
        style_name = (para.style.name if para.style else "Normal") or "Normal"
        text       = para.text.strip()
        display    = text[:100] if text else "(vacío)"
        in_table   = is_in_table(para)

        # ── 0. Contexto de tabla (solo si el usuario lo activa) ─────
        if apply_table_context and in_table and style_name in NON_TABLE_BODY_STYLES:
            confidence = 0.92 if style_name == "Normal" else 0.78
            results.append(_make_result(
                i, display, style_name,
                suggested="Table Paragraph",
                confidence=confidence,
                status="review",
                note=f"en tabla, estilo '{style_name}' → Table Paragraph",
            ))
            continue

        # ── 1. Heading mal aplicado (sin bold) ───────────────────────
        if style_name in DEMOTABLE_HEADINGS:
            misapplied, conf, reason = is_misapplied_heading(para)
            if misapplied:
                results.append(_make_result(
                    i, display, style_name,
                    suggested="Normal",
                    confidence=conf,
                    status="review",
                    note=reason,
                ))
                continue
            results.append(_make_result(
                i, display, style_name,
                suggested=style_name,
                confidence=1.0,
                status="no_change",
            ))
            continue

        # ── 2. List Paragraph con bold → posible Heading 2 ───────────
        if style_name == "List Paragraph" and not in_table:
            likely_h2, conf_h2, reason_h2 = is_likely_subheading(para)
            if likely_h2:
                results.append(_make_result(
                    i, display, style_name,
                    suggested="Heading 2",
                    confidence=conf_h2,
                    status="review",
                    note=reason_h2,
                ))
                continue

        # ── 3. Estilo reconocido por Oxygen ──────────────────────────
        if style_name in known_style_names:
            results.append(_make_result(
                i, display, style_name,
                suggested=style_name,
                confidence=1.0,
                status="no_change",
            ))
            continue

        # ── 4. Regla guardada ─────────────────────────────────────────
        if style_name in rules_dict:
            target, rule_conf = rules_dict[style_name]
            if target in HEADING_TARGETS:
                likely, heur_conf, reason = is_likely_heading(para)
                if likely:
                    results.append(_make_result(
                        i, display, style_name,
                        suggested=target,
                        confidence=heur_conf,
                        status="auto" if heur_conf >= 0.85 else "review",
                        note=reason,
                    ))
                else:
                    results.append(_make_result(
                        i, display, style_name,
                        suggested=style_name,
                        confidence=0.0,
                        status="no_change",
                    ))
            else:
                results.append(_make_result(
                    i, display, style_name,
                    suggested=target,
                    confidence=rule_conf,
                    status="auto" if rule_conf >= 0.80 else "review",
                ))
            continue

        # ── 5. Sin regla — ¿parece heading? ──────────────────────────
        likely, heur_conf, reason = is_likely_heading(para)
        if likely:
            results.append(_make_result(
                i, display, style_name,
                suggested="Heading 1",
                confidence=heur_conf,
                status="review",
                note=reason,
            ))
            continue

        # ── 6. Estilo tolerable ───────────────────────────────────────
        if style_name in TOLERATED_STYLES:
            results.append(_make_result(
                i, display, style_name,
                suggested=style_name,
                confidence=1.0,
                status="no_change",
            ))
            continue

        # ── 7. Sin regla — intentar coincidencia por similitud de nombre ──
        match, score, reason = find_closest_style(style_name, known_style_names)
        if match:
            results.append(_make_result(
                i, display, style_name,
                suggested=match,
                confidence=score,
                status="review",
                note=f"sugerencia por similitud de nombre ({reason})",
            ))
            continue

        # ── 8. Sin mapeo ──────────────────────────────────────────────
        results.append(_make_result(
            i, display, style_name,
            suggested=style_name,
            confidence=0.0,
            status="unknown",
        ))

    return results


def _make_result(idx, display, original_style,
                 suggested, confidence, status, note=""):
    return {
        "idx":             idx,
        "text":            display,
        "original_style":  original_style,
        "suggested_style": suggested,
        "final_style":     suggested,
        "confidence_pct":  round(confidence * 100),
        "status":          status,
        "status_label":    STATUS_LABELS[status],
        "note":            note,
    }


# ── Aplicar decisiones ────────────────────────────────────────────────────────

def _get_or_create_paragraph_style(doc, style_name):
    """
    Devuelve el estilo de párrafo con ese nombre. Si el documento no lo tiene
    definido, lo crea.

    Esto es necesario porque Word solo guarda en el catálogo de estilos del
    documento (styles.xml) los estilos que se han usado al menos una vez.
    Documentos generados por otras herramientas (conversión desde HTML, CMS,
    plantillas automatizadas) a menudo solo incluyen 'Normal' y ningún estilo
    de Heading, aunque el nombre sea estándar de Word. Sin este parche, intentar
    aplicar 'Heading 1' a un párrafo en ese documento fallaría con KeyError.
    """
    try:
        style = doc.styles[style_name]
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            return style
    except KeyError:
        pass

    new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    try:
        new_style.base_style = doc.styles["Normal"]
    except KeyError:
        pass

    # Formato razonable para que el estilo tenga sentido si se abre en Word
    name_lower = style_name.lower()
    if name_lower.startswith("heading "):
        try:
            level = int(style_name.split()[-1])
        except ValueError:
            level = 1
        new_style.font.bold = True
        new_style.font.size = Pt(max(11, 16 - level))
        new_style.paragraph_format.keep_with_next = True
    elif style_name == "Title":
        new_style.font.bold = True
        new_style.font.size = Pt(22)
    elif style_name == "Quote":
        new_style.font.italic = True

    return new_style


def apply_paragraph_decisions(input_path, output_path, decisions):
    """
    Aplica estilos finales párrafo a párrafo.
    Devuelve (n_cambios, n_errores).
    """
    doc     = Document(input_path)
    paras   = _get_all_paras(doc)
    changes = 0
    errors  = 0

    for d in decisions:
        idx            = d["idx"]
        final_style    = d["final_style"]
        original_style = d["original_style"]

        if final_style in (original_style, "(sin cambio)") or idx >= len(paras):
            continue
        try:
            target_style = _get_or_create_paragraph_style(doc, final_style)
            paras[idx].style = target_style
            changes += 1
        except Exception:
            errors += 1

    doc.save(output_path)
    return changes, errors


def extract_styles_from_docx(file_path):
    """Resumen de estilos usados (función legacy)."""
    try:
        doc = Document(file_path)
    except PackageNotFoundError:
        raise ValueError("El archivo no es un .docx válido.")
    styles = {}
    def _add(para):
        name = (para.style.name if para.style else "Normal") or "Normal"
        if name not in styles:
            styles[name] = {"element_type": "p", "count": 0, "sample": "",
                            "heading_candidates": []}
        styles[name]["count"] += 1
        if not styles[name]["sample"] and para.text.strip():
            styles[name]["sample"] = para.text.strip()[:120]
        likely, conf, reason = is_likely_heading(para)
        if likely:
            styles[name]["heading_candidates"].append(
                {"text": para.text.strip()[:80], "confidence": conf, "reason": reason})
    for para in doc.paragraphs:
        _add(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _add(para)
    return styles
