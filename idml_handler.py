"""
idml_handler.py
Adaptador de InDesign (.idml) para Style Mapper.

FILOSOFÍA DE DISEÑO: este módulo NO decide qué es un heading ni corrige
nada. Su único trabajo es "transcribir" el contenido y los nombres de
estilo de un archivo .idml a un documento .docx sintético — como si
alguien hubiera escrito ese mismo contenido en Word usando esos mismos
nombres de estilo (aunque no existan de forma nativa en Word).

Una vez transcrito, el documento entra en el pipeline existente
(analyze_document, apply_paragraph_decisions, reglas, heurísticas) SIN
NINGÚN CAMBIO — todo lo que ya funciona para DOCX funciona igual aquí.

Limitación conocida: el orden de lectura entre distintas "Stories" de
InDesign se aproxima usando el atributo StoryList de designmap.xml (el
orden de creación/enlace de los textos), que no siempre coincide al
100% con el orden visual de lectura en maquetaciones complejas
(columnas, textos flotantes). Es una aproximación razonable para
detectar y corregir estilos — no pretende reconstruir la maquetación.
"""

import html
import re
import zipfile

from docx import Document
from lxml import etree

from style_processor import _get_or_create_paragraph_style

# ── Normalización de nombres de estilo de InDesign ────────────────────────────
# El catálogo "moderno" observado (ver sesión de análisis) organiza los
# headings en una convención por tipo de topic DITA: "H1 - Topic",
# "H2 - Concept", "H2 - Reference", "H2 - Task", además de "Heading 1/2..."
# genéricos. Confirmado con el equipo: esa distinción de tipo de topic NO
# se preserva en Word — todo colapsa al nivel numérico genérico.
_HEADING_PATTERN = re.compile(r"^H(\d)\s*-\s*.+$")
_GENERIC_HEADING_PATTERN = re.compile(r"^Heading\s*(\d)$", re.IGNORECASE)


def _normalize_style_name(raw_name: str) -> str:
    """
    Normaliza mecánicamente los nombres de heading de InDesign a la forma
    genérica de Word ("H2 - Concept" → "Heading 2"). Es una normalización
    puramente sintáctica de una convención de nomenclatura conocida — no
    es un juicio sobre si el párrafo "es" o "no es" un heading; eso lo
    sigue decidiendo style_processor.py como con cualquier otro formato.
    """
    m = _HEADING_PATTERN.match(raw_name)
    if m:
        return f"Heading {m.group(1)}"
    m = _GENERIC_HEADING_PATTERN.match(raw_name)
    if m:
        return f"Heading {m.group(1)}"
    return raw_name


def _style_ref_to_leaf_name(style_ref: str) -> str:
    """
    Convierte una referencia de estilo de IDML (ej.
    'ParagraphStyle/New Paragraph Styles%3aHeadings%3aH1 - Topic') en el
    nombre de estilo "hoja" (el último segmento del path de grupos),
    que es el nombre real que el usuario ve en InDesign.
    """
    # Quitar el prefijo de tipo ("ParagraphStyle/")
    ref = style_ref.split("/", 1)[-1] if "/" in style_ref else style_ref
    ref = html.unescape(ref)
    # Los grupos de estilos usan ':' codificado como %3a
    ref = ref.replace("%3a", ":").replace("%3A", ":")

    if ref in ("$ID/[No paragraph style]", "[No paragraph style]"):
        return "Normal"
    if ref in ("$ID/NormalParagraphStyle", "NormalParagraphStyle"):
        return "Normal"

    leaf = ref.split(":")[-1].strip()
    return leaf if leaf else "Normal"


def _unescape_idml_text(text: str) -> str:
    """Convierte entidades XML y caracteres especiales de IDML a texto plano."""
    text = html.unescape(text)
    # <Br/> ya se filtra antes de llegar aquí; aquí solo limpiamos espacios raros
    return text.replace("\u2028", " ").replace("\u2029", " ")


def extract_idml_paragraphs(idml_path: str) -> list:
    """
    Lee un .idml y devuelve una lista de dicts en un orden de lectura
    aproximado: {"text": str, "style_name": str, "is_bold": bool}.

    Usa un parser XML real (lxml) en lugar de expresiones regulares,
    porque IDML anida <ParagraphStyleRange> dentro de tablas (una tabla
    dentro de un párrafo, con sus propios párrafos dentro de cada
    celda) — un enfoque basado en regex pierde la mayoría del contenido
    en documentos con tablas, al confundir el cierre de una etiqueta
    anidada con el de la etiqueta exterior.

    style_name ya viene normalizado para los headings de la convención
    "H<N> - <tipo>" / "Heading <N>" → "Heading <N>". Cualquier otro
    nombre (Body, Table Body, nombres antiguos tipo "103.5 Heading 1",
    nombres personalizados de marketing, etc.) se devuelve tal cual —
    el pipeline existente decide qué hacer con ellos, igual que con
    cualquier estilo desconocido de un .docx.
    """
    results = []

    with zipfile.ZipFile(idml_path) as z:
        try:
            designmap = z.read("designmap.xml").decode("utf-8", errors="ignore")
        except KeyError:
            designmap = ""

        story_match = re.search(r'StoryList="([^"]*)"', designmap)
        story_ids = story_match.group(1).split() if story_match else []

        all_story_files = sorted(
            n for n in z.namelist() if n.startswith("Stories/") and n.endswith(".xml")
        )
        if story_ids:
            story_files = [f"Stories/Story_{sid}.xml" for sid in story_ids]
            story_files = [f for f in story_files if f in z.namelist()]
            story_files += [f for f in all_story_files if f not in story_files]
        else:
            story_files = all_story_files

        for story_path in story_files:
            try:
                raw = z.read(story_path)
            except KeyError:
                continue

            try:
                root = etree.fromstring(raw)
            except etree.XMLSyntaxError:
                continue

            # .iter() recorre TODO el árbol (incluidas las ParagraphStyleRange
            # anidadas dentro de tablas) en orden de documento — a diferencia
            # de un regex, esto entiende el anidamiento correctamente.
            for para_el in root.iter("ParagraphStyleRange"):
                style_ref = para_el.get("AppliedParagraphStyle", "")
                style_name = _normalize_style_name(_style_ref_to_leaf_name(style_ref))

                text_parts = []
                any_bold_with_text = False

                # Solo los CharacterStyleRange DIRECTAMENTE dentro de este
                # párrafo — no los de una tabla anidada dentro de él, que
                # se procesarán como sus propios párrafos por separado
                # cuando .iter() llegue a ellos más adelante.
                for char_el in para_el.findall("CharacterStyleRange"):
                    font_style = char_el.get("FontStyle", "")
                    fragment_parts = [
                        (content_el.text or "")
                        for content_el in char_el.findall("Content")
                    ]
                    fragment = "".join(fragment_parts)
                    if fragment.strip():
                        text_parts.append(fragment)
                        if "bold" in font_style.lower():
                            any_bold_with_text = True

                text = "".join(text_parts).strip()
                if not text:
                    continue

                results.append({
                    "text": text,
                    "style_name": style_name,
                    "is_bold": any_bold_with_text,
                })

    return results


def build_synthetic_docx(idml_path: str, output_path: str) -> int:
    """
    Extrae el contenido de un .idml y genera un .docx sintético en
    output_path: mismo texto, con el estilo de párrafo correspondiente
    (creado al vuelo si no existe) y negrita aplicada donde InDesign la
    tenía. Ese .docx es indistinguible, para el resto del pipeline, de
    uno subido directamente por el usuario.

    Devuelve el número de párrafos transcritos.
    """
    paragraphs = extract_idml_paragraphs(idml_path)

    doc = Document()
    # Vaciar el párrafo vacío por defecto del documento en blanco
    if doc.paragraphs and not doc.paragraphs[0].text:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    style_cache = {}

    for p in paragraphs:
        style_name = p["style_name"]
        if style_name not in style_cache:
            style_cache[style_name] = _get_or_create_paragraph_style(doc, style_name)
        style_obj = style_cache[style_name]

        para = doc.add_paragraph(style=style_obj)
        run = para.add_run(p["text"])
        if p["is_bold"]:
            run.bold = True

    doc.save(output_path)
    return len(paragraphs)
